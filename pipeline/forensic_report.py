"""Day 3 — Forensic appendix docx renderer.

Produces the Step 6 forensic audit document. The same renderer handles both
report types (Pre-Application / Post-Registration), branching on
ForensicReport.report_type for sections that differ (executive summary tone,
'Recommended Specification' vs 'Recommended Actions', final recommendation).

Layout (10 sections):
  1.  Cover panel
  2.  Introduction
  3.  Executive Summary
  4.  Overall Risk Assessment
  5.  Methodology
  6.  Scoring Table (sorted descending)
  7.  Trademark-by-Trademark Detail (cards)
  8.  Recommended Specification (Pre) / Recommended Actions (Post)
  9.  Final Professional Recommendation
  10. Sign-off (Alex Pugh)
"""
from __future__ import annotations
from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .report_builder import (
    set_cell_bg, set_cell_border, risk_fill, add_hyperlink, run,
    add_para, add_heading,
)
from .brand_tokens import (
    BRAND_PINK, BRAND_NAVY, BRAND_SLATE, BRAND_LIGHT_SLATE,
    BRAND_BODY, BRAND_WHITE, BRAND_FONT,
    BRAND_TAGLINE, LOGO_PATH_FULL, logo_exists,
)
from .forensic_narrative import ForensicReport, ReportType


# ---------- shared helpers ----------

def _h1(doc, text):
    """Forensic-specific heading style — slightly different colour palette
    from the monitoring report's add_heading."""
    return add_heading(doc, text, level=1)


def _h2(doc, text):
    return add_heading(doc, text, level=2)


def _h3(doc, text):
    return add_heading(doc, text, level=3)


def _make_kv_table(doc, rows, *, key_width=2.0, val_width=5.0, key_fill=BRAND_NAVY):
    """Standard two-column key/value table used in cover, exec summary etc."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        for ci, txt in enumerate([k, v]):
            cell = t.rows[i].cells[ci]
            cell.text = ''
            set_cell_border(cell)
            cell.width = Inches(key_width if ci == 0 else val_width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if ci == 0:
                set_cell_bg(cell, key_fill)
                run(p, str(txt), bold=True, color=BRAND_WHITE, size=10)
            else:
                lines = str(txt).split('\n')
                run(p, lines[0], size=10)
                for line in lines[1:]:
                    p2 = cell.add_paragraph()
                    p2.paragraph_format.space_after = Pt(0)
                    run(p2, line, size=10)
    return t


def _safe_get(d, *keys, default=''):
    """Walk a possibly-nested dict by key path, returning default if missing."""
    cur = d
    for k in keys:
        if not isinstance(cur, dict) or k not in cur:
            return default
        cur = cur[k]
    return cur if cur is not None else default


# ---------- 1. Cover panel ----------

def _render_cover(doc, report: ForensicReport, order_meta: dict):
    """Title + order meta table at the start of the document."""
    # TMH logo at top, just like the initial report
    if logo_exists(LOGO_PATH_FULL):
        p_logo = doc.add_paragraph()
        p_logo.paragraph_format.space_after = Pt(6)
        p_logo.add_run().add_picture(LOGO_PATH_FULL, width=Inches(2.4))

    title_label = ('Forensic Audit \u2014 Pre-Application Trademark Risk'
                   if report.report_type == ReportType.PRE_APPLICATION
                   else 'Forensic Audit \u2014 Post-Registration Infringement Review')

    # Mark Type indicator \u2014 Word Mark vs Image Mark \u2014 so the cover matches
    # the initial report's convention.
    word_or_image = (order_meta.get('word_or_image') or '').strip().lower()
    is_image_report = 'image' in word_or_image or 'logo' in word_or_image or 'figurative' in word_or_image
    mark_type_label = 'Image Mark' if is_image_report else 'Word Mark'

    add_para(doc, order_meta.get('client_name', ''), bold=True,
             color=BRAND_NAVY, size=16, space_after=2)
    add_para(doc, f'{title_label} \u2014 {mark_type_label}',
             bold=True, color=BRAND_NAVY, size=18, space_after=10)

    brand_ref = order_meta.get('brand_reference') or report.client_brand.get('brand_reference') or '—'
    report_ref = order_meta.get('report_reference') or '—'

    cover_rows = [
        ['Prepared for', order_meta.get('client_name', '')],
        ['Brand Reference', brand_ref],
        ['Report Reference', report_ref],
        ['Mark', order_meta.get('mark_label', '')],
        ['Trademark Classes', order_meta.get('classes', '')],
        ['Designated Countries', order_meta.get('countries', '')],
        ['Date of Audit', date.today().strftime('%d %B %Y')],
        ['Report Type', 'Pre-Application' if report.report_type == ReportType.PRE_APPLICATION else 'Post-Registration'],
        ['Records Audited', f'{len(report.records)} cited senior records'],
        ['Verification Source', f"Signa API ({sum(1 for r in report.records if r.verified)} of {len(report.records)} records verified)"],
        ['Prepared By', order_meta.get('prepared_by', '')],
        ['Account Manager', order_meta.get('account_manager', '')],
    ]
    _make_kv_table(doc, cover_rows)


# ---------- 2. Introduction ----------

def _render_introduction(doc, report: ForensicReport):
    _h1(doc, '1. Introduction')
    if report.report_type == ReportType.PRE_APPLICATION:
        add_para(doc, (
            'This appendix supplements the Braudit monitoring report by providing '
            'forensic verification of every cited senior trademark, scored against '
            'four criteria (classes & terms overlap, mark type, age & proof of use, '
            'region & jurisdiction). The report is specifically focused on the '
            'likelihood of objection or refusal that the client\u2019s proposed '
            'trademark application may face from the examining authority or third '
            'parties during the publication period.'
        ), size=11)
        add_para(doc, (
            'Every record cited in this appendix has been verified against the '
            'Signa unified trademark API, which mirrors the underlying registers '
            '(USPTO, EUIPO, WIPO and others) and is synced daily with the source '
            'offices. Where verification was unsuccessful for a specific record '
            '(API timeout or no matching entry), the record is flagged explicitly '
            'with a note requiring manual verification.'
        ), size=11)
    else:
        add_para(doc, (
            'This appendix supplements the Braudit monitoring report by providing '
            'forensic verification of every potentially-infringing third-party '
            'trademark, scored against four criteria (classes & terms overlap, '
            'mark type, age & proof of use, region & jurisdiction). The report is '
            'specifically focused on potential infringement of the client\u2019s '
            'registered rights and the strength of the case for enforcement action.'
        ), size=11)
        add_para(doc, (
            'Every record cited in this appendix has been verified against the '
            'Signa unified trademark API, which mirrors the underlying registers '
            '(USPTO, EUIPO, WIPO and others) and is synced daily with the source '
            'offices. Where verification was unsuccessful, the record is flagged '
            'explicitly with a note requiring manual verification.'
        ), size=11)


# ---------- 3. Executive Summary ----------

def _render_executive_summary(doc, report: ForensicReport):
    _h1(doc, '2. Executive Summary')
    summary_text = _safe_get(report.summary, 'executive_summary',
                              default='[Executive summary unavailable.]')
    for paragraph in str(summary_text).split('\n\n'):
        if paragraph.strip():
            add_para(doc, paragraph.strip(), size=11)


# ---------- 4. Overall Risk Assessment ----------

def _render_risk_assessment(doc, report: ForensicReport):
    _h1(doc, '3. Overall Risk Assessment')
    from collections import Counter
    band_counts = Counter(s.risk_band for s in report.scores)
    total = len(report.scores)
    avg_score = (sum(s.total for s in report.scores) / total) if total else 0.0

    if report.report_type == ReportType.PRE_APPLICATION:
        viability_score = _safe_get(report.summary, 'viability_score', default='—')
        viability_label = _safe_get(report.summary, 'viability_label', default='—')
        rows = [
            ['Average forensic score across cited records', f'{avg_score:.1f} / 10'],
            ['Highest individual record score', f'{max((s.total for s in report.scores), default=0)} / 10'],
            ['Records at Very High risk', f"{band_counts.get('Very High', 0)} of {total}"],
            ['Records at High risk', f"{band_counts.get('High', 0)} of {total}"],
            ['Records at Medium risk', f"{band_counts.get('Medium', 0)} of {total}"],
            ['Records at Low or Negligible risk', f"{band_counts.get('Low', 0) + band_counts.get('Negligible', 0)} of {total}"],
            ['Filing Viability Score', f'{viability_score} / 10'],
            ['Viability Assessment', str(viability_label)],
        ]
    else:
        enforcement_priority = _safe_get(report.summary, 'enforcement_priority',
                                          default='Monitor Only')
        rows = [
            ['Average forensic score across cited records', f'{avg_score:.1f} / 10'],
            ['Highest individual record score', f'{max((s.total for s in report.scores), default=0)} / 10'],
            ['Records at Critical / Very High enforcement priority', f"{band_counts.get('Very High', 0)} of {total}"],
            ['Records at High enforcement priority', f"{band_counts.get('High', 0)} of {total}"],
            ['Records at Medium enforcement priority', f"{band_counts.get('Medium', 0)} of {total}"],
            ['Records at Low / Monitor-only priority', f"{band_counts.get('Low', 0) + band_counts.get('Negligible', 0)} of {total}"],
            ['Headline Enforcement Priority', str(enforcement_priority)],
        ]

    _make_kv_table(doc, rows)


# ---------- 5. Methodology ----------

def _render_methodology(doc, report: ForensicReport):
    _h1(doc, '4. Methodology')
    add_para(doc, (
        'Each cited record was verified against the Signa unified trademark '
        'API. The verified record was then scored on a deterministic 0\u201310 '
        'rubric across four criteria:'
    ), size=11)
    add_para(doc, '\u2022  Classes and terms overlap with the client\u2019s filing (0\u20133)', size=11, space_after=2)
    add_para(doc, '\u2022  Type of mark \u2014 word marks have broader scope than figurative (0\u20132)', size=11, space_after=2)
    add_para(doc, '\u2022  Age and proof of use \u2014 live registered rights score highest (0\u20133)', size=11, space_after=2)
    add_para(doc, '\u2022  Region and jurisdiction \u2014 same-jurisdiction rights weigh more (0\u20132)', size=11)
    add_para(doc, (
        'Per-record forensic commentary was then generated by Claude Sonnet 4.6 '
        'using the Run Card methodology, framed for the report type. Verbatim '
        'register data (recital, dates, owner, status) is preserved exactly as '
        'returned by Signa; only the forensic commentary paragraph is '
        'LLM-generated.'
    ), size=11)


# ---------- 6. Scoring Table ----------

def _render_scoring_table(doc, report: ForensicReport):
    _h1(doc, '5. Scoring Table')
    add_para(doc, (
        'All cited records sorted by descending forensic score. Risk colour '
        'codes \u2014 Very High (dark red), High (orange), Medium (amber), Low '
        '(green), Negligible (grey).'
    ), size=10)

    # Pair records with scores, sort by score descending
    pairs = sorted(
        zip(report.records, report.scores),
        key=lambda p: -p[1].total,
    )

    headers = ['Office', 'App #', 'Mark Text', 'Owner', 'Status', 'Classes', 'Score', 'Risk']
    # App # widened from 0.85 → 1.05 so UK numbers (UK00003076168 = 13 chars)
    # fit. Recovered from Owner (1.7 → 1.5) and Classes (0.6 → 0.5). Sums to 7.0".
    widths = [0.55, 1.05, 1.40, 1.50, 0.85, 0.50, 0.50, 0.65]
    t = doc.add_table(rows=1 + len(pairs), cols=len(headers))
    t.autofit = False

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = ''
        set_cell_bg(c, BRAND_NAVY)
        set_cell_border(c)
        c.width = Inches(widths[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run(p, h, bold=True, color=BRAND_WHITE, size=8)

    # Body rows
    for ri, (record, score) in enumerate(pairs, start=1):
        row = t.rows[ri]
        vals = [
            record.office or '—',
            record.app_number or '—',
            record.mark_text or '—',
            record.owner_name or '—',
            (record.status or '—').title(),
            ', '.join(str(c) for c in record.nice_classes) or '—',
            f'{score.total}',
            score.risk_band,
        ]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = ''
            set_cell_border(cell)
            cell.width = Inches(widths[ci])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if ci == 7:  # Risk column — colour-fill
                set_cell_bg(cell, risk_fill(score.risk_band))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if ci == 1 and record.source_url:
                add_hyperlink(p, record.source_url, str(val), font_size=8)
            else:
                run(p, str(val), size=8,
                    bold=(ci in (6, 7)))


# ---------- 7. Trademark-by-Trademark Detail (cards) ----------

def _render_record_cards(doc, report: ForensicReport):
    _h1(doc, '6. Trademark-by-Trademark Detail')
    add_para(doc, (
        'Every cited record below, in the order shown in the scoring table. Each '
        'card contains verbatim Signa-verified data, followed by forensic '
        'commentary written for this audit\u2019s report type.'
    ), size=10)

    pairs = sorted(
        enumerate(zip(report.records, report.scores)),
        key=lambda i_rs: -i_rs[1][1].total,
    )

    for original_idx, (record, score) in pairs:
        record_id = f'record_{original_idx}'
        _render_one_card(doc, record, score,
                         report.commentaries.get(record_id, ''),
                         report.report_type)
        # Spacer paragraph between cards
        add_para(doc, '', size=4, space_after=4)


def _render_one_card(doc, record, score, commentary: str, report_type: ReportType):
    # Card title — mark text + app number + register hyperlink
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, f'{record.mark_text or "(unknown mark)"}  \u2014  ',
        bold=True, color=BRAND_NAVY, size=13)
    label = f'{record.office} {record.app_number}'
    if record.source_url:
        add_hyperlink(p, record.source_url, label, font_size=11)
    else:
        run(p, label, size=11)

    # Risk pill
    pill = doc.add_paragraph()
    pill.paragraph_format.space_after = Pt(6)
    pill_text_run = pill.add_run(f' {score.risk_band} \u2014 score {score.total}/10 ')
    pill_text_run.font.size = Pt(9)
    pill_text_run.font.bold = True
    pill_text_run.font.color.rgb = RGBColor.from_string(
        'FFFFFF' if score.risk_band in ('Very High', 'High') else '000000'
    )
    # Background fill for the pill via cell isn't possible on a run, so we
    # just embed the colour via the bg of a one-cell table instead:
    # (Keeping it simple here — the scoring table already shows the band.)

    # Two-column field grid
    fields = [
        ['Office', record.office or '—'],
        ['Application Number', record.app_number or '—'],
        ['Registration Number', record.registration_number or '—'],
        ['Mark Type', f'{record.mark_type or "—"} ({record.mark_legal_category or "—"})'],
        ['Status', f"{record.status or '—'} \u2014 {record.status_stage or '—'}"],
        ['Filing / Registration', f"Filed {record.filing_date or '—'}  |  Reg {record.registration_date or '—'}"],
        ['Expiry / Renewal Due', f"{record.expiry_date or '—'}  |  {record.renewal_due_date or '—'}"],
        ['Owner', record.owner_name or '—'],
        ['Nice Classes', ', '.join(str(c) for c in record.nice_classes) or '—'],
        ['Verified', 'Yes \u2014 ' + record.verification_note
                       if record.verified else 'NO \u2014 ' + record.verification_note],
    ]
    _make_kv_table(doc, fields, key_width=1.8, val_width=5.7)

    # Goods and services (verbatim)
    if record.goods_services:
        add_para(doc, 'Verbatim Goods & Services:', bold=True, size=10,
                 space_after=2)
        for gs in record.goods_services:
            add_para(doc, '  \u2022  ' + gs, size=10, space_after=2)

    # Forensic score breakdown
    add_para(doc, 'Forensic Score Breakdown:', bold=True, size=10, space_after=2)
    add_para(doc,
             f'  Classes & terms {score.classes_terms}/3  \u00b7  '
             f'Mark type {score.mark_type}/2  \u00b7  '
             f'Age & proof of use {score.age_proof_of_use}/3  \u00b7  '
             f'Region {score.region}/2',
             size=10, space_after=4)
    for n in score.explanation:
        add_para(doc, '  \u2022  ' + n, size=10, space_after=2)

    # LLM-generated commentary
    label_text = ('Likelihood of Objection \u2014 Forensic Commentary'
                  if report_type == ReportType.PRE_APPLICATION
                  else 'Enforcement Priority \u2014 Forensic Commentary')
    add_para(doc, label_text, bold=True, size=10, space_after=2,
             color=BRAND_NAVY)
    if commentary:
        for para in commentary.split('\n\n'):
            if para.strip():
                add_para(doc, para.strip(), size=10)
    else:
        add_para(doc, '[Commentary unavailable for this record.]',
                 italic=True, color=BRAND_LIGHT_SLATE, size=10)


# ---------- 8. Recommended Specification (Pre) / Recommended Actions (Post) ----------

def _render_pre_recommended_spec(doc, report: ForensicReport):
    _h1(doc, '7. Recommended Specification')
    add_para(doc, (
        'Draft Nice-class wording designed to walk past the top-cited senior '
        'records identified in this audit. Negative limitations included where '
        'helpful to distinguish your goods from the highest-scoring conflicts. '
        'Specialist in-jurisdiction counsel should refine this language before '
        'filing.'
    ), size=10)
    spec = _safe_get(report.summary, 'recommended_specification', default={})
    if isinstance(spec, dict) and spec:
        for cls_key, text in spec.items():
            _h3(doc, f'Class {cls_key}')
            add_para(doc, str(text), size=11)
    else:
        add_para(doc, '[No recommended specification generated.]',
                 italic=True, color=BRAND_LIGHT_SLATE, size=10)


def _render_post_recommended_actions(doc, report: ForensicReport):
    _h1(doc, '7. Recommended Actions')
    add_para(doc, (
        'Per-record enforcement recommendations. Actions are prioritised by '
        'forensic score and commercial impact. Specialist counsel should be '
        'engaged before issuing any cease-and-desist letter or opposition.'
    ), size=10)
    actions = _safe_get(report.summary, 'recommended_actions', default=[])
    if isinstance(actions, list) and actions:
        # Action table: record / action / rationale
        headers = ['Record', 'Recommended Action', 'Rationale']
        widths = [2.0, 1.7, 3.3]
        t = doc.add_table(rows=1 + len(actions), cols=3)
        t.autofit = False
        # Header
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ''
            set_cell_bg(c, BRAND_NAVY)
            set_cell_border(c)
            c.width = Inches(widths[i])
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run(p, h, bold=True, color=BRAND_WHITE, size=10)
        # Rows
        for ai, action_obj in enumerate(actions, start=1):
            record_id = (action_obj.get('record_id') if isinstance(action_obj, dict) else '') or ''
            action = (action_obj.get('action') if isinstance(action_obj, dict) else '') or ''
            rationale = (action_obj.get('rationale') if isinstance(action_obj, dict) else '') or ''
            # Look up the record so we can show owner / mark
            label = record_id
            try:
                if record_id.startswith('record_'):
                    idx = int(record_id.split('_', 1)[1])
                    rec = report.records[idx]
                    label = f'{rec.mark_text}\n{rec.owner_name}\n{rec.office} {rec.app_number}'
            except Exception:
                pass
            cells = [label, action, rationale]
            for ci, val in enumerate(cells):
                cell = t.rows[ai].cells[ci]
                cell.text = ''
                set_cell_border(cell)
                cell.width = Inches(widths[ci])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                lines = str(val).split('\n')
                run(p, lines[0], size=10,
                    bold=(ci == 1))
                for line in lines[1:]:
                    p2 = cell.add_paragraph()
                    p2.paragraph_format.space_after = Pt(0)
                    run(p2, line, size=10)
    else:
        add_para(doc, '[No recommended actions generated.]',
                 italic=True, color=BRAND_LIGHT_SLATE, size=10)


# ---------- 9. Final Recommendation ----------

def _render_final_recommendation(doc, report: ForensicReport):
    _h1(doc, '8. Final Professional Recommendation')
    text = _safe_get(report.summary, 'final_recommendation',
                      default='[Final recommendation unavailable.]')
    for paragraph in str(text).split('\n\n'):
        if paragraph.strip():
            add_para(doc, paragraph.strip(), size=11)


# ---------- 10. Sign-off ----------

def _render_signoff(doc, order_meta: dict):
    _h1(doc, '9. Sign-off')
    prepared_by = order_meta.get('prepared_by', 'Alex Pugh')
    account_manager = order_meta.get('account_manager', '')
    add_para(doc, f'Prepared by:  {prepared_by}', bold=True, size=11, space_after=2)
    add_para(doc, 'Trademark Intelligence & Brand Protection Specialist',
             size=11, space_after=2)
    add_para(doc, 'Braudit Trademark Intelligence Division',
             italic=True, color=BRAND_NAVY, size=11, space_after=8)
    if account_manager:
        add_para(doc, f'Account Manager:  {account_manager}', size=11, space_after=2)
    add_para(doc, f'Date:  {date.today().strftime("%d %B %Y")}',
             size=11, space_after=8)
    add_para(doc, (
        'This appendix is prepared as a forensic supplement to the parent '
        'Braudit monitoring/representation report. It is not legal advice and '
        'does not substitute for specialist in-jurisdiction trademark counsel. '
        'Braudit recommends specialist counsel handling for any subsequent '
        'filing, opposition or enforcement action arising from this audit.'
    ), italic=True, color=BRAND_LIGHT_SLATE, size=10)


# ---------- top-level entrypoint ----------

def build_forensic_appendix(report: ForensicReport, order_meta: dict) -> bytes:
    """Render a ForensicReport into a Word docx and return the bytes."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    _render_cover(doc, report, order_meta)
    _render_introduction(doc, report)
    _render_executive_summary(doc, report)
    _render_risk_assessment(doc, report)
    _render_methodology(doc, report)
    _render_scoring_table(doc, report)
    _render_record_cards(doc, report)
    if report.report_type == ReportType.PRE_APPLICATION:
        _render_pre_recommended_spec(doc, report)
    else:
        _render_post_recommended_actions(doc, report)
    _render_final_recommendation(doc, report)
    _render_signoff(doc, order_meta)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
