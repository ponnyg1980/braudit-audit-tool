"""Step 5: build the Braudit-style Word document from the filtered/scored data.

Branding is sourced from `brand_tokens.py` (TMH colours, logo, tagline) so
every report this module produces stays in sync with the live brand.
"""
from __future__ import annotations
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from datetime import date
from io import BytesIO
from .nice_classes import NICE_HEADINGS, parse_classes
from .brand_tokens import (
    BRAND_PINK, BRAND_NAVY, BRAND_SLATE, BRAND_LIGHT_SLATE,
    BRAND_BODY, BRAND_WHITE, BRAND_BORDER, BRAND_FONT,
    BRAND_TAGLINE, LOGO_PATH_FULL, logo_exists,
    RISK_VERY_HIGH, RISK_HIGH, RISK_MEDIUM, RISK_LOW, RISK_NEGLIGIBLE, RISK_CLIENT_LIKELY,
    RISK_THRESHOLDS, USABLE_PAGE_WIDTH_IN,
)


# ---------- styling helpers ----------

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)


def set_cell_border(cell, color='BFBFBF', size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        borders.append(b)
    tc_pr.append(borders)


def risk_fill(risk):
    """Map a risk-band label to the cell fill colour. Bands are semantic
    (red/orange/amber/green/grey) — not brand colours — so they stay stable
    across rebrands. Source of truth: brand_tokens.RISK_*."""
    r = (risk or '').lower()
    # BR-IMG-003: 'Client Likely' is a non-severity category and must be
    # checked BEFORE the substring-based severity checks so 'Client Likely'
    # doesn't accidentally match the 'low' substring inside 'likely'.
    if 'client likely' in r: return RISK_CLIENT_LIKELY
    if 'very high' in r: return RISK_VERY_HIGH
    if 'high' in r: return RISK_HIGH
    if 'medium' in r: return RISK_MEDIUM
    if 'low' in r: return RISK_LOW
    if 'negligible' in r: return RISK_NEGLIGIBLE
    return BRAND_WHITE


def tm_url(office: str, app_num: str) -> str:
    """Return the official-register URL for a trademark, based on the office code.

    Supports US (USPTO TSDR), UK (UKIPO), EU (EUIPO eSearch) and WIPO/Madrid
    (WIPO Brand Database). Returns '' for unknown offices so callers can fall
    back to plain text instead of a broken link.
    """
    if not app_num:
        return ''
    o = (office or '').upper().strip()
    n = str(app_num).strip()
    if o in ('US', 'USPTO'):
        return f'https://tsdr.uspto.gov/statusview/sn{n}'
    if o in ('UK', 'UKIPO', 'GB'):
        return f'https://trademarks.ipo.gov.uk/ipo-tmcase/page/Results/1/{n}'
    if o in ('EU', 'EUIPO', 'EM'):
        return f'https://euipo.europa.eu/eSearch/#details/trademarks/{n}'
    if o in ('WIPO', 'MADRID', 'IR', 'WO'):
        return f'https://www3.wipo.int/branddb/en/showData.jsp?ID={n}'
    return ''


def add_hyperlink(paragraph, url, text, color=BRAND_PINK, underline=True, font_size=9, bold=False):
    """Inserts an external hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), BRAND_FONT); rFonts.set(qn('w:hAnsi'), BRAND_FONT)
    r_pr.append(rFonts)

    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(font_size * 2)); r_pr.append(sz)
    color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), color); r_pr.append(color_el)
    if underline:
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); r_pr.append(u)
    if bold:
        b = OxmlElement('w:b'); r_pr.append(b)
    new_run.append(r_pr)
    t = OxmlElement('w:t'); t.text = text; t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def run(p, text, *, bold=False, italic=False, color=BRAND_BODY, size=11, font=BRAND_FONT):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def add_para(doc, text='', *, bold=False, italic=False, color=BRAND_BODY, size=11, align=None, space_after=4):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run(p, text, bold=bold, italic=italic, color=color, size=size)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 16, 2: 13, 3: 11}
    # TMH palette: H1/H2 in navy, H3 in slate
    colors = {1: BRAND_NAVY, 2: BRAND_NAVY, 3: BRAND_SLATE}
    run(p, text, bold=True, color=colors.get(level, BRAND_BODY), size=sizes.get(level, 12))
    if level == 1:
        # Bottom border in pink for visual punch
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '8'); bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), BRAND_PINK)
        pBdr.append(bottom); pPr.append(pBdr)
    return p


def add_table(doc, col_widths_in, header_row, body_rows, *, header_fill=BRAND_NAVY, header_color=BRAND_WHITE,
              risk_col_index=None, hyperlink_col_indexes=None, font_size=9):
    """Generic styled table.
    body_rows: list[list[str]] or list[list[(text, opts_dict)]]
    hyperlink_col_indexes: dict {col_index: lambda row -> url} to render that cell as a link
    risk_col_index: int — colour that cell by risk text
    """
    table = doc.add_table(rows=1 + len(body_rows), cols=len(header_row))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # set column widths
    for col_idx, w_in in enumerate(col_widths_in):
        for cell in table.columns[col_idx].cells:
            cell.width = Inches(w_in)

    # header
    hdr = table.rows[0]
    for i, txt in enumerate(header_row):
        c = hdr.cells[i]
        c.text = ''
        set_cell_bg(c, header_fill)
        set_cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run(p, txt, bold=True, color=header_color, size=font_size)

    # body
    hyperlink_col_indexes = hyperlink_col_indexes or {}
    for r_idx, row in enumerate(body_rows):
        tr = table.rows[1 + r_idx]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ''
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            # risk cell fill
            if risk_col_index is not None and c_idx == risk_col_index:
                set_cell_bg(cell, risk_fill(str(val)))

            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)

            # NEW: image cell support \u2014 if val is a dict with 'image_bytes',
            # render the JPEG/PNG inline at the requested width. Falls back to
            # empty if there are no bytes.
            if isinstance(val, dict) and 'image_bytes' in val:
                img_bytes = val.get('image_bytes')
                img_w = float(val.get('width_in', 0.55))
                if img_bytes:
                    try:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell.paragraphs[0].add_run().add_picture(BytesIO(img_bytes), width=Inches(img_w))
                    except Exception as exc:
                        run(p, f'[img err: {exc.__class__.__name__}]', size=font_size, color='C00000')
                # If no image bytes, leave the cell visually empty
                continue

            # NEW: multi-link cell support \u2014 if val is a list of (text, url)
            # tuples, render each as its own hyperlink on its own paragraph.
            if isinstance(val, list) and val and isinstance(val[0], tuple):
                first_text, first_url = val[0]
                if first_url:
                    add_hyperlink(p, first_url, str(first_text), font_size=font_size)
                else:
                    run(p, str(first_text), size=font_size)
                for txt, url in val[1:]:
                    p_next = cell.add_paragraph()
                    p_next.paragraph_format.space_after = Pt(0)
                    if url:
                        add_hyperlink(p_next, url, str(txt), font_size=font_size)
                    else:
                        run(p_next, str(txt), size=font_size)
                continue

            if c_idx in hyperlink_col_indexes:
                url = hyperlink_col_indexes[c_idx](row)
                if url:
                    add_hyperlink(p, url, str(val), font_size=font_size)
                else:
                    run(p, str(val), size=font_size)
            else:
                txt = str(val) if val is not None else ''
                # Split on \n into multiple paragraphs
                lines = txt.split('\n')
                run(p, lines[0], size=font_size,
                    bold=(c_idx == risk_col_index))
                for line in lines[1:]:
                    p2 = cell.add_paragraph()
                    p2.paragraph_format.space_after = Pt(0)
                    run(p2, line, size=font_size)

    return table


# ---------- main builder ----------

def build_step5_report(*, order_meta: dict,
                       trademarks_live: list[dict],
                       trademarks_dead: list[dict],
                       companies: list[dict],
                       google: list[dict],
                       domains: list[dict],
                       social: list[dict],
                       raw_counts: dict) -> bytes:
    """Generate the Step 5 Braudit-style monitoring report as docx bytes."""
    doc = Document()

    # Set default font + margins
    style = doc.styles['Normal']
    style.font.name = BRAND_FONT
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ---- TMH logo (top of cover) ----
    if logo_exists(LOGO_PATH_FULL):
        p_logo = doc.add_paragraph()
        p_logo.paragraph_format.space_after = Pt(6)
        p_logo.add_run().add_picture(LOGO_PATH_FULL, width=Inches(2.4))

    # ---- Header / Title ----
    # Tri-state report-type detection from order_meta.word_or_image:
    #   "Word Only"               -> Word Mark Report
    #   "Logo or Figurative Mark" -> Image Mark Report
    #   "Combined Word and Logo"  -> Combined Word + Logo Report
    word_or_image = (order_meta.get('word_or_image') or '').strip().lower()
    is_combined_report = 'combined' in word_or_image
    is_image_report = (not is_combined_report) and (
        'image' in word_or_image or 'logo' in word_or_image or 'figurative' in word_or_image
    )
    # Reports that need the image cover treatment (logo embed + Vienna codes)
    # are both Image and Combined.
    show_image_cover = is_combined_report or is_image_report
    if is_combined_report:
        mark_type_label = 'Combined Word + Logo Report'
    elif is_image_report:
        mark_type_label = 'Image Mark Report'
    else:
        mark_type_label = 'Word Mark Report'

    add_para(doc, order_meta.get('client_name', ''), bold=True, color=BRAND_NAVY, size=16, space_after=2)
    add_para(doc, f'Monitoring or Representation Report — {mark_type_label}',
             bold=True, color=BRAND_NAVY, size=18, space_after=10)

    # For Image Mark and Combined reports: embed the logo/figurative image
    # inline so the cover panel visibly shows what we're searching for.
    # Image bytes come from the Order Form parser (extract_order_metadata
    # captures B30/B31). For Combined reports, the word search criteria are
    # already shown in Section 2 — only the image side needs cover treatment.
    if show_image_cover:
        img1 = order_meta.get('image_1_bytes')
        img2 = order_meta.get('image_2_bytes')
        if img1 or img2:
            add_para(doc, 'Image(s) being searched:', bold=True, color=BRAND_SLATE,
                     size=10, space_after=4)
            for blob, label in [(img1, 'Image Mark 1'), (img2, 'Image Mark 2')]:
                if blob:
                    p_img = doc.add_paragraph()
                    p_img.paragraph_format.space_after = Pt(2)
                    try:
                        p_img.add_run().add_picture(BytesIO(blob), width=Inches(1.6))
                        add_para(doc, label, italic=True, color=BRAND_LIGHT_SLATE,
                                 size=9, space_after=6)
                    except Exception:
                        # Don't crash the whole report on a bad image
                        add_para(doc, f'[{label}: image could not be rendered]',
                                 italic=True, color=BRAND_LIGHT_SLATE, size=9)
            # Vienna classifications block — structured list of (code, description)
            # from the Order Form's D31:D40 + E31:E40. Fall back to the legacy
            # vienna_classes string if the structured list is empty.
            vienna_list = order_meta.get('vienna_classifications') or []
            vienna_str = (order_meta.get('vienna_classes') or '').strip()
            if vienna_list:
                add_para(doc, 'Vienna classifications (client image):',
                         bold=True, color=BRAND_SLATE, size=10, space_after=2)
                v_rows = [[v.get('code', '') or '—', v.get('description', '') or '—']
                          for v in vienna_list]
                add_table(doc, [1.5, 5.5],
                          ['Classification', 'Description'],
                          v_rows, font_size=10)
                add_para(doc, '', size=4, space_after=6)
            elif vienna_str:
                add_para(doc, f'Vienna classifications (client image): {vienna_str}',
                         bold=True, color=BRAND_SLATE, size=10, space_after=10)

    # ---- Cover / Order Detail Table ----
    # Build the "Trademark Classes" value with UKIPO standardised definitions
    class_nums = parse_classes(order_meta.get('classes', ''))
    if class_nums:
        class_lines = []
        for n in class_nums:
            heading = NICE_HEADINGS.get(n, '')
            if heading:
                class_lines.append(f"Class {n} \u2014 {heading}")
            else:
                class_lines.append(f"Class {n}")
        classes_display = "\n".join(class_lines)
    else:
        classes_display = order_meta.get('classes', '')

    # Build the "Specific Terms" value from the client's actual G&S per class
    specific_terms = order_meta.get('specific_terms') or {}
    if specific_terms:
        st_lines = []
        for n in class_nums or sorted(specific_terms.keys()):
            terms = specific_terms.get(n) or specific_terms.get(str(n)) or ''
            if terms:
                st_lines.append(f"Class {n}: {terms}")
        specific_terms_display = "\n".join(st_lines) if st_lines else '\u2014'
    else:
        specific_terms_display = '\u2014'

    cover_rows = [
        ['Prepared for', order_meta.get('client_name', '')],
        ['Brand Reference', order_meta.get('brand_reference', '') or '\u2014'],
        ['Report Reference', order_meta.get('report_reference', '') or '\u2014'],
        ['Client Contact', f"{order_meta.get('client_first','')} {order_meta.get('client_last','')}".strip()],
        ['Client Email', order_meta.get('client_email', '')],
        ['Account Manager', order_meta.get('account_manager', '')],
        ['Report Prepared By', order_meta.get('prepared_by', '')],
        ['Date of Search', order_meta.get('search_date', '')],
        ['Type of Search', order_meta.get('search_type', 'Word')],
        ['Mark', order_meta.get('mark_label', '')],
        ['Trademark Classes', classes_display],
        ['Specific Terms', specific_terms_display],
        ['SIC Code', order_meta.get('sic', '')],
        ['Nature of Business', order_meta.get('nature', '')],
        ['Designated Countries', order_meta.get('countries', '')],
        ['Filtering Rules', order_meta.get('filtering_rules', '')],
    ]
    t = doc.add_table(rows=len(cover_rows), cols=2)
    t.autofit = False
    # Widths sum to USABLE_PAGE_WIDTH_IN (7.0") so the table does NOT
    # overflow the right margin (the 10 Jun 2026 fix). Was 2.2 + 5.3 = 7.5.
    for i, (k, v) in enumerate(cover_rows):
        for ci, txt in enumerate([k, v]):
            cell = t.rows[i].cells[ci]
            cell.text = ''
            set_cell_border(cell)
            cell.width = Inches(2.0 if ci == 0 else 5.0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if ci == 0:
                set_cell_bg(cell, BRAND_NAVY)
                run(p, txt, bold=True, color=BRAND_WHITE, size=10)
            else:
                # Multi-line cell support \u2014 each \n becomes a paragraph
                lines = str(txt).split('\n')
                run(p, lines[0], size=10)
                for line in lines[1:]:
                    p2 = cell.add_paragraph()
                    p2.paragraph_format.space_after = Pt(0)
                    run(p2, line, size=10)

    # ---- Section 1 ----
    add_heading(doc, '1. Report Overview', level=1)
    add_para(doc, 'This report provides a monitoring update based on searches carried out in the last 30 days. It compiles findings from multiple sources, including search engines, company registries, domain name databases, social media platforms, and trademark registries. The purpose is to identify any existing use of the name or logo, potential infringements, or similar trademarks that could conflict with the client\u2019s branding.', size=11)
    add_para(doc, 'This report is for guidance only and does not replace expert trademark advice or professional legal advice.', size=11)

    add_heading(doc, 'Results Overview', level=2)
    overview_rows = [
        ['Search Engine (Google)', str(raw_counts.get('google_raw', 0)), str(len(google))],
        ['Companies House (UK)', str(raw_counts.get('companies_raw', 0)), str(len(companies))],
        ['Domains', str(raw_counts.get('domains_raw', 0)), str(len(domains))],
        ['Social Media', str(raw_counts.get('social_raw', 0)), str(len(social))],
        ['Trademark Registers \u2014 Live', str(raw_counts.get('tm_live_raw', 0)), str(len(trademarks_live))],
        ['Trademark Registers \u2014 Dead (Negligible)', str(raw_counts.get('tm_dead_raw', 0)), str(len(trademarks_dead))],
    ]
    # Widths sum to 7.0" \u2014 was 3.5 + 1.8 + 2.2 = 7.5" (overflowed margin).
    add_table(doc, [3.4, 1.6, 2.0],
              ['Platform', 'Total Results', 'Flagged On This Report'],
              overview_rows, font_size=10)

    # ---- Section 2 ----
    add_heading(doc, '2. Search Criteria', level=1)

    # Word and domain search criteria are now multi-row (up to 5 of each)
    # with one of four types (Exact Match / Starts With / Contains / Similar
    # To) plus optional remarks. Pull from the structured lists if present,
    # otherwise fall back to the back-compat single-value fields so reports
    # generated before the multi-row update still render sensibly.
    word_searches = order_meta.get('word_searches') or []
    domain_searches = order_meta.get('domain_searches') or []

    # Word searches table
    add_para(doc, 'Word Search', bold=True)
    if word_searches:
        w_body = [[w.get('type', ''), w.get('phrase', ''), w.get('remarks', '')]
                  for w in word_searches]
        add_table(doc, [1.5, 3.5, 2.0],
                  ['Search Type', 'Phrase', 'Remarks'],
                  w_body, font_size=10)
    else:
        # Back-compat path
        if order_meta.get('exact'):
            add_para(doc, f"Exact Match: {order_meta.get('exact','')}", space_after=2)
        if order_meta.get('similar'):
            add_para(doc, f"Starts With: {order_meta.get('similar','')}", space_after=2)
        if not order_meta.get('exact') and not order_meta.get('similar'):
            add_para(doc, 'No word search criteria captured.',
                     italic=True, color=BRAND_LIGHT_SLATE, space_after=2)
    add_para(doc, '', space_after=4)

    # Domain searches table
    add_para(doc, 'Domain Search', bold=True)
    if domain_searches:
        d_body = [[d.get('type', ''), d.get('phrase', ''), d.get('remarks', '')]
                  for d in domain_searches]
        add_table(doc, [1.5, 3.5, 2.0],
                  ['Search Type', 'Phrase', 'Remarks'],
                  d_body, font_size=10)
    else:
        add_para(doc, 'No domain search criteria captured.',
                 italic=True, color=BRAND_LIGHT_SLATE, space_after=2)
    add_para(doc, '', space_after=8)

    add_para(doc, 'Exclusion Rules Applied', bold=True)
    add_para(doc, 'Trademarks: marks where the search root is not the leading word were dropped (e.g. BAY STEALTH, SUPER STEALTH, SPOTLIGHT STEALTH). Records that did not touch any of the client\u2019s classes were dropped. Dead/ended records were retained but tagged Negligible Risk.', size=10)
    add_para(doc, 'Companies House: records were retained only where the registered SIC included the client\u2019s SIC.', size=10)
    add_para(doc, 'Domains / Google / Social: the raw scrape contained one aggregated row per mark variant, so no further filtering was required.', size=10)

    # ---- Section 3 ----
    doc.add_page_break()
    add_heading(doc, '3. Search Results, Risk Grading and Analysis', level=1)

    # 3a Google
    add_heading(doc, '3a. Google Search Results', level=2)
    add_para(doc, 'Word Searches use Google\u2019s matching criteria. Image Searches use Google Lens to identify similar images online. Image results render the actual image found; word results render the keyword.', size=10)
    if google:
        # The Keyword column renders EITHER an inline image (when the
        # original Google sheet had a Picture-in-Cell rich-value image at
        # column A) OR the text keyword. add_table's image-cell support
        # recognises a dict carrying 'image_bytes'.
        def _kw_cell(g: dict):
            img = g.get('image_bytes')
            if img:
                return {'image_bytes': img, 'width_in': 1.4}
            return g.get('keyword', '')

        # Dropped the legacy "Score" column \u2014 it carried a magic number (30.04
        # / 44.35) from the original STEALTH-LED scoring code that has no
        # documented meaning to clients. Risk band carries the signal; the
        # Link column now gets the recovered width so URLs stop clipping
        # off the right margin. Widths sum to 7.0" (was 7.8" and overflowed).
        g_body = [[_kw_cell(g), g['risk'], g['link']] for g in google]
        add_table(doc, [1.6, 1.4, 4.0],
                  ['Keyword / Image', 'Risk', 'Link'],
                  g_body, risk_col_index=1,
                  hyperlink_col_indexes={2: lambda row: row[2]},
                  font_size=10)
    else:
        add_para(doc, 'No Google results flagged.', italic=True, color=BRAND_LIGHT_SLATE)

    # 3b Companies
    add_heading(doc, '3b. Companies House UK Search Results', level=2)
    add_para(doc, 'Company name search using UK Companies House. Records retained only where the SIC included client SIC. Company name and company number link directly to the Companies House register.', size=10)
    if companies:
        # Build a {number: link} map so both Name and Co. No. cells can use it
        co_link_by_number = {c['number']: c.get('link') or f"https://find-and-update.company-information.service.gov.uk/company/{c['number']}" for c in companies}
        c_body = [[c['mark'], c['status'], c['number'], c['sic'], c['risk']] for c in companies]
        # Widths sum to 7.0" — was 7.8" and overflowed.
        add_table(doc, [2.3, 0.9, 1.3, 1.4, 1.1],
                  ['Registered Company', 'Status', 'Co. No.', 'SIC', 'Risk'],
                  c_body, risk_col_index=4,
                  hyperlink_col_indexes={
                      0: lambda row: co_link_by_number.get(row[2], ''),  # mark name -> CH page
                      2: lambda row: co_link_by_number.get(row[2], ''),  # co. no -> CH page
                  },
                  font_size=10)
    else:
        add_para(doc, 'No matching companies flagged.', italic=True, color=BRAND_LIGHT_SLATE)

    # 3c Domains
    add_heading(doc, '3c. Domain Name Search Results', level=2)
    add_para(doc, 'Domain searches using variations of the client\u2019s mark across .com, .net, .co.uk, .co and .uk.', size=10)
    if domains:
        # Build domain cells as list[(text, url)] so each URL is a clickable link.
        # Dropped the "(score)" suffix — that number was a STEALTH-LED magic
        # value (40.76 / 63) with no documented scale. Risk band carries the signal.
        d_body = [
            [d['mark_text'],
             [(u, u) for u in d['urls']],
             d['risk']]
            for d in domains
        ]
        # Widths sum to 7.0" — was 7.8" and overflowed.
        add_table(doc, [1.9, 3.9, 1.2],
                  ['Mark Variant', 'Domains', 'Risk'], d_body,
                  risk_col_index=2, font_size=10)
    else:
        add_para(doc, 'No domain results flagged.', italic=True, color=BRAND_LIGHT_SLATE)

    # 3d Social
    add_heading(doc, '3d. Social Media Search Results', level=2)
    add_para(doc, 'Social media searches across Facebook, Instagram, LinkedIn, TikTok, YouTube and X. Each platform name is a clickable link to that account.', size=10)
    if social:
        s_body = []
        for s in social:
            # Build platform cells as list[(label, url)] so each platform name
            # is a clickable hyperlink straight to that profile/page.
            plats = [(k, v) for k, v in s['platforms'].items() if v]
            # Magic-number (40.76 / 63) suffix dropped — same rationale as 3c.
            s_body.append([s['mark_text'], plats, s['risk']])
        # Widths sum to 7.0" — was 7.8" and overflowed.
        add_table(doc, [1.9, 3.9, 1.2],
                  ['Mark Variant', 'Platforms', 'Risk'], s_body,
                  risk_col_index=2, font_size=9)
    else:
        add_para(doc, 'No social results flagged.', italic=True, color=BRAND_LIGHT_SLATE)

    # Helper: build the per-row Class cell value as "11 \u2014 heading\n12 \u2014 heading"
    def _class_cell(cls_str):
        nums = parse_classes(cls_str)
        if not nums:
            return str(cls_str or '')
        lines = []
        for n in nums:
            h = NICE_HEADINGS.get(n, '')
            if h:
                lines.append(f"{n} \u2014 {h}")
            else:
                lines.append(str(n))
        return "\n".join(lines)

    # 3e Trademarks Summary
    # BR-IMG-005 (10 Jun 2026): 3e is the at-a-glance summary table — one
    # row per cited mark with just the class NUMBERS (no UKIPO definitions).
    # 3f below is the detailed click-to-collapse Heading-3 view. Live and
    # Dead records are combined here, sorted by risk grade (Client Likely
    # > High > Medium > Low > Negligible at bottom).
    doc.add_page_break()
    add_heading(doc, '3e. Trademark Search Results — Summary', level=1)
    add_para(doc, 'At-a-glance summary of every cited mark. Records are sorted by risk grade — Client Likely first, then High Risk, Medium, Low, Negligible (Dead) at the bottom. Scroll down to section 3f for the detail view per mark, including class definitions, goods/services and the embedded logo image.', size=10)

    # Scoring legend so the operator and the client can interpret the Risk
    # column without having to read the methodology section. Thresholds are
    # mirrored from brand_tokens.RISK_THRESHOLDS, which mirrors filters.py.
    legend_rows = [[band, threshold] for band, threshold in RISK_THRESHOLDS]
    add_para(doc, 'Scoring legend (initial-review score is out of 13):',
             bold=True, color=BRAND_SLATE, size=10, space_after=2)
    add_table(doc, [1.5, 5.5],
              ['Risk Band', 'Score Range'],
              legend_rows, risk_col_index=0, font_size=10)
    add_para(doc, '', size=4, space_after=4)

    # Shorten the "Stylized characters" Mark Type label so it fits on one line.
    def _short_type(t):
        s = str(t or '').strip()
        if 'stylized' in s.lower():
            return 'Stylized'
        return s

    # Class column for the summary table: numbers ONLY, no UKIPO definitions.
    # User feedback (10 Jun 2026): the at-a-glance table is meant to be
    # scannable — full definitions live in section 3f's detailed view.
    def _class_numbers_only(classes_str: str) -> str:
        if not classes_str:
            return ''
        # Already comma-separated like '9, 11, 35' — pass through unchanged.
        return str(classes_str).strip()

    # Visual-similarity cell (BR-IMG-002). Logo rows show CLIP cosine or a
    # pHash distance marker. Word rows show em dash.
    def _visual_cell(t):
        if t.get('threat_type') != 'Logo':
            return '—'
        cos = t.get('visual_clip_cosine', -1.0)
        if cos is not None and cos >= 0:
            return f'{cos:.2f}'
        d = t.get('visual_phash_distance', -1)
        if d is not None and d >= 0:
            prefix = 'd<=' if d <= 5 else 'd'
            return f'{prefix}{d}'
        return '—'

    def _img_cell(t):
        return {'image_bytes': t.get('image_bytes'), 'width_in': 0.55}

    def _app_link(row):
        return tm_url(row[0], row[1])

    # Summary-table column layout. Class column shrinks dramatically (just
    # numbers) so other columns get more breathing room. Widths sum to 7.0".
    if is_combined_report:
        # 11 cols, total: 0.40+0.95+0.95+0.50+0.50+0.45+0.85+0.70+0.50+0.50+0.70 = 7.00
        tm_widths  = [0.40,    0.95,    0.95,        0.50,    0.50,     0.45,   0.85,    0.70,     0.50,     0.50,     0.70]
        tm_headers = ['Office','App #', 'Mark Text', 'Image', 'Classes','Type', 'Owner', 'Status', 'Threat', 'Visual', 'Risk']
    else:
        # 9 cols, total: 0.40+1.00+1.10+0.55+0.55+0.55+1.10+0.85+0.90 = 7.00
        tm_widths  = [0.40,    1.00,    1.10,        0.55,    0.55,     0.55,   1.10,    0.85,     0.90]
        tm_headers = ['Office','App #', 'Mark Text', 'Image', 'Classes','Type', 'Owner', 'Status', 'Risk']

    def _summary_row(t):
        base = [t['office'], t['app'], t['mark'], _img_cell(t),
                _class_numbers_only(t['classes']), _short_type(t['type']),
                t['owner'], t['status']]
        if is_combined_report:
            return base + [t.get('threat_type', 'Word'), _visual_cell(t), t['risk']]
        return base + [t['risk']]

    risk_idx = 10 if is_combined_report else 8

    # Combine live + dead into one summary table; both are already sorted.
    all_marks = list(trademarks_live) + list(trademarks_dead)
    if all_marks:
        body = [_summary_row(t) for t in all_marks]
        add_table(doc, tm_widths, tm_headers, body,
                  risk_col_index=risk_idx,
                  hyperlink_col_indexes={1: _app_link},
                  font_size=7)
        # Combined-report-only visual signal legend.
        if is_combined_report:
            add_para(doc, '', size=4, space_after=2)
            add_para(doc,
                     'Visual column: CLIP cosine similarity (0.00-1.00) between the client logo and each cited mark logo, '
                     'or pHash distance when CLIP was skipped (d<= near-duplicate, d>30 visually unrelated). '
                     '>= 0.85 = visually identical (High if Live + same industry, else Medium). '
                     '0.65-0.84 = mild similarity (Medium retained). '
                     '< 0.65 = visually unrelated (demoted to Low).',
                     italic=True, color=BRAND_LIGHT_SLATE, size=8, space_after=4)
    else:
        add_para(doc, 'No trademarks flagged.', italic=True, color=BRAND_LIGHT_SLATE)

    # ---------------------------------------------------------------------
    # 3f Trademark Detail — click-to-collapse Heading-3 blocks per cited mark.
    # ---------------------------------------------------------------------
    doc.add_page_break()
    add_heading(doc, '3f. Trademark Search Results — Detail', level=1)
    add_para(doc, 'Full detail per cited mark. Each entry below is collapsible in Word — click the triangle next to the heading line to fold or unfold its detail. The entries appear in the same risk-sorted order as section 3e above.', size=10)

    def _summary_visual(t: dict) -> str:
        """Compact visual-signal label for the H3 line. Empty when no signal."""
        cos = t.get('visual_clip_cosine', -1.0)
        if cos is not None and cos >= 0:
            return f'cos {cos:.2f}'
        d = t.get('visual_phash_distance', -1)
        if d is not None and d >= 0:
            prefix = 'd<=' if d <= 5 else 'd'
            return f'{prefix}{d}'
        return ''

    def _cited_mark_block(t: dict) -> None:
        """One cited mark = Heading-3 summary line + collapsible body."""
        office = (t.get('office') or '').strip()
        app = (t.get('app') or '').strip()
        owner = (t.get('owner') or '').strip()
        risk = (t.get('risk') or '').strip()
        threat = t.get('threat_type', '') if is_combined_report else ''

        office_app = f'{office} {app}' if office and app else (app or office or '-')
        parts = [office_app, owner or '-']
        if threat:
            parts.append(threat)
        parts.append(risk)
        visual_str = _summary_visual(t)
        if visual_str and (not threat or threat == 'Logo'):
            parts.append(visual_str)
        summary = '   |   '.join(parts)

        # Heading 3 paragraph (Word renders the click-to-collapse triangle).
        p = doc.add_paragraph(style='Heading 3')
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        rh = p.add_run(summary)
        rh.font.name = BRAND_FONT; rh.font.size = Pt(10); rh.font.bold = True
        rh.font.color.rgb = RGBColor.from_string(BRAND_NAVY)

        # ----- Body (collapses when user clicks the triangle) -----
        mark = (t.get('mark') or '').strip()
        img = t.get('image_bytes')
        if img or mark:
            body = doc.add_paragraph()
            body.paragraph_format.space_after = Pt(2)
            body.paragraph_format.left_indent = Inches(0.25)
            if img:
                try:
                    body.add_run().add_picture(BytesIO(img), width=Inches(0.7))
                    body.add_run('   ')
                except Exception:
                    pass
            if mark:
                rb = body.add_run('Mark: ')
                rb.font.bold = True; rb.font.name = BRAND_FONT; rb.font.size = Pt(10)
                rb2 = body.add_run(mark)
                rb2.font.name = BRAND_FONT; rb2.font.size = Pt(10)

        meta_bits = []
        if t.get('type'):   meta_bits.append(f'Type: {t["type"]}')
        if t.get('status'): meta_bits.append(f'Status: {t["status"]}')
        if t.get('filing'): meta_bits.append(f'Filing: {t["filing"]}')
        if meta_bits:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            p2.paragraph_format.left_indent = Inches(0.25)
            r = p2.add_run('   |   '.join(meta_bits))
            r.font.name = BRAND_FONT; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(BRAND_BODY)

        classes = (t.get('classes') or '').strip()
        if classes:
            class_text = _class_cell(classes)  # FULL UKIPO defs here in the detail
            p3 = doc.add_paragraph()
            p3.paragraph_format.space_after = Pt(2)
            p3.paragraph_format.left_indent = Inches(0.25)
            r = p3.add_run('Classes: ')
            r.font.bold = True; r.font.name = BRAND_FONT; r.font.size = Pt(9)
            r2 = p3.add_run(class_text)
            r2.font.name = BRAND_FONT; r2.font.size = Pt(9)

        goods = (t.get('goods') or '').strip()
        if goods:
            if len(goods) > 350:
                goods = goods[:350] + '...'
            p4 = doc.add_paragraph()
            p4.paragraph_format.space_after = Pt(2)
            p4.paragraph_format.left_indent = Inches(0.25)
            r = p4.add_run('Goods/Services: ')
            r.font.bold = True; r.font.name = BRAND_FONT; r.font.size = Pt(9)
            r2 = p4.add_run(goods)
            r2.font.name = BRAND_FONT; r2.font.size = Pt(9)

        decision = (t.get('visual_decision') or '').strip()
        if decision and decision != 'skipped':
            bits = []
            cos = t.get('visual_clip_cosine', -1.0)
            if cos is not None and cos >= 0:
                bits.append(f'CLIP cosine {cos:.3f}')
            ph = t.get('visual_phash_distance', -1)
            if ph is not None and ph >= 0:
                bits.append(f'pHash distance {ph}/64')
            rationale = f'Visual similarity: {decision}'
            if bits:
                rationale += f' ({", ".join(bits)})'
            p5 = doc.add_paragraph()
            p5.paragraph_format.space_after = Pt(4)
            p5.paragraph_format.left_indent = Inches(0.25)
            r = p5.add_run(rationale)
            r.font.italic = True; r.font.name = BRAND_FONT; r.font.size = Pt(8)
            r.font.color.rgb = RGBColor.from_string(BRAND_LIGHT_SLATE)

        url = tm_url(office, app)
        if url:
            p6 = doc.add_paragraph()
            p6.paragraph_format.space_after = Pt(8)
            p6.paragraph_format.left_indent = Inches(0.25)
            r = p6.add_run('View on register: ')
            r.font.name = BRAND_FONT; r.font.size = Pt(8)
            r.font.color.rgb = RGBColor.from_string(BRAND_SLATE)
            add_hyperlink(p6, url, url, color=BRAND_PINK,
                          underline=True, font_size=8)

    if all_marks:
        for _cm in all_marks:
            _cited_mark_block(_cm)
    else:
        add_para(doc, 'No trademarks flagged.', italic=True, color=BRAND_LIGHT_SLATE)

    # ---- Footer message ----
    doc.add_page_break()
    add_heading(doc, 'A message from our founder', level=1)
    add_para(doc, 'Thank you for taking the time to review this report.')
    add_para(doc, 'We have been helping businesses protect their brands since 2008, and if anything within this document needs clarification or gives you cause for concern, please do not hesitate to contact us.')
    add_para(doc, 'If you require full representation or only assistance with a particular stage, our team is ready to support you.')
    add_para(doc, 'We continually look to improve our service, so any feedback you can provide is always welcome.')
    add_para(doc, ' ')
    add_para(doc, 'Jonathan Paton', bold=True, color=BRAND_NAVY)
    add_para(doc, 'Founder and Director', color=BRAND_SLATE)
    add_para(doc, 'The Trademark Helpline', color=BRAND_SLATE)
    add_para(doc, '')
    # Brand sign-off — tagline in TMH pink, centred.
    add_para(doc, BRAND_TAGLINE, bold=True, color=BRAND_PINK, size=12,
             align='center', space_after=0)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
